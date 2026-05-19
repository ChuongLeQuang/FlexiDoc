"""
EN: Core service for processing and validating Word templates.
VI: Dịch vụ lõi để xử lý và xác thực mẫu Word.
"""

import io
import re
import docx
import jinja2
from docxtpl import DocxTemplate
from typing import Dict, List
from src.exceptions.custom_errors import (
    InvalidFileFormatError,
    TemplateSyntaxError,
    ZeroVariableError,
)


def extract_variables(template_bytes: bytes) -> Dict[str, List[str]]:
    """
    EN: Extract all single and table variables from the Word template, preserving their order of appearance.
    VI: Trích xuất tất cả các biến đơn và biến bảng biểu từ mẫu Word, bảo toàn thứ tự xuất hiện của chúng.
    """
    try:
        doc = docx.Document(io.BytesIO(template_bytes))
    except Exception:
        raise InvalidFileFormatError("Định dạng file Word không hợp lệ hoặc bị hỏng.")

    full_text = []
    for p in doc.paragraphs:
        full_text.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    full_text.append(p.text)

    combined_text = "\n".join(full_text)

    # Tìm tất cả các thẻ {{ ... }}
    matches = re.findall(r"\{\{\s*(.*?)\s*\}\}", combined_text)

    single_vars = {}
    table_vars = {}
    current_table = None

    for match in matches:
        match = match.strip()
        if match.startswith("table_start:") or match.startswith("list_start:"):
            current_table = match.split(":", 1)[1].strip()
            if current_table not in table_vars:
                table_vars[current_table] = {}
        elif match == "table_end" or match == "list_end":
            current_table = None
        else:
            if current_table:
                # Nếu biến nằm trong bảng, ví dụ {{item.name}} hoặc {{name}}, lưu lại 'name'
                var_name = match[5:] if match.startswith("item.") else match
                table_vars[current_table][var_name] = None
            else:
                single_vars[match] = None

    if not single_vars and not table_vars:
        raise ZeroVariableError("Không tìm thấy biến nào trong mẫu Word.")

    return {
        "single_vars": list(single_vars.keys()),
        "table_vars": {k: list(v.keys()) for k, v in table_vars.items()},
    }


def preprocess_template(template_bytes: bytes) -> bytes:
    """
    EN: Convert custom simple table/list syntax to Jinja2 syntax. This process might remove rich text formatting from the lines containing the tags.
    VI: Chuyển đổi cú pháp bảng biểu/danh sách đơn giản thành cú pháp của Jinja2. Quá trình này có thể làm mất định dạng rich text trên các dòng chứa thẻ.
    """
    try:
        doc = docx.Document(io.BytesIO(template_bytes))

        def replace_in_paragraphs(paragraphs, is_table=False):
            tag = "tr" if is_table else "p"
            for p in paragraphs:
                if any(
                    keyword in p.text
                    for keyword in [
                        "table_start",
                        "table_end",
                        "list_start",
                        "list_end",
                    ]
                ):
                    new_text = p.text
                    new_text = re.sub(
                        r"\{\{\s*(?:table_start|list_start)\s*:\s*([^}]+?)\s*\}\}",
                        f"{{%{tag} for item in \\1 %}}",
                        new_text,
                    )
                    new_text = re.sub(
                        r"\{\{\s*(?:table_end|list_end)\s*\}\}",
                        f"{{%{tag} endfor %}}",
                        new_text,
                    )
                    p.text = new_text

        replace_in_paragraphs(doc.paragraphs, is_table=False)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    replace_in_paragraphs(cell.paragraphs, is_table=True)

        out = io.BytesIO()
        doc.save(out)
        return out.getvalue()
    except Exception as e:
        raise InvalidFileFormatError(f"Lỗi trong quá trình tiền xử lý file Word: {e}")


def validate_syntax(template_bytes: bytes) -> None:
    """
    EN: Perform a dry-run to validate Jinja2 syntax in the template.
    VI: Chạy thử để kiểm tra tính hợp lệ của cú pháp Jinja2.
    """
    preprocessed_bytes = preprocess_template(template_bytes)
    try:
        tpl = DocxTemplate(io.BytesIO(preprocessed_bytes))
        tpl.render({})
    except jinja2.exceptions.TemplateSyntaxError as e:
        raise TemplateSyntaxError(f"Lỗi cú pháp Jinja2 tại dòng: {str(e)}")
    except Exception as e:
        raise TemplateSyntaxError(f"Lỗi cú pháp mẫu Word: {str(e)}")
