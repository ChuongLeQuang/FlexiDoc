"""
EN: Script to generate sample Word and Excel files for manual UI testing.
VI: Kịch bản sinh file Word và Excel mẫu để test thực tế trên giao diện.
"""

import os
import docx
import openpyxl


def generate_word():
    doc = docx.Document()
    doc.add_heading("PHỤ LỤC HỢP ĐỒNG BÀN GIAO THIẾT BỊ", 0)

    doc.add_paragraph("Họ và tên nhân viên: {{ ho_ten }}")
    doc.add_paragraph("Ngày sinh: {{ ngay_sinh }}")
    doc.add_paragraph("Mức lương cơ bản: {{ muc_luong }}")
    doc.add_paragraph("Bằng chữ: {{ muc_luong_chu }}")

    doc.add_heading("DANH SÁCH BÀN GIAO", level=2)

    # Tạo bảng biểu 3 cột, có kẻ khung
    table = doc.add_table(rows=4, cols=3)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "STT"
    table.rows[0].cells[1].text = "Tên thiết bị"
    table.rows[0].cells[2].text = "Số lượng"

    # Row 1: Thẻ bắt đầu vòng lặp
    table.rows[1].cells[0].text = "{{ table_start:ds_thiet_bi }}"

    # Row 2: Dòng dữ liệu thực tế
    table.rows[2].cells[0].text = "{{ item.stt }}"
    table.rows[2].cells[1].text = "{{ item.ten_thiet_bi }}"
    table.rows[2].cells[2].text = "{{ item.so_luong }}"

    # Row 3: Thẻ kết thúc vòng lặp
    table.rows[3].cells[0].text = "{{ table_end }}"

    os.makedirs("samples", exist_ok=True)
    try:
        doc.save("samples/Mau_Hop_Dong.docx")
        print("✅ Đã tạo file: samples/Mau_Hop_Dong.docx")
    except PermissionError:
        print(
            "❌ LỖI: Không thể lưu. Vui lòng đóng file 'samples/Mau_Hop_Dong.docx' nếu đang mở trong Word và chạy lại."
        )


def generate_excel():
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "NhanSu"

    # Dòng 1: Tiêu đề
    headers = [
        "Mã NV",
        "Họ và tên",
        "Ngày sinh",
        "Mức lương",
        "STT",
        "Tên thiết bị",
        "Số lượng",
    ]
    ws.append(headers)

    # Dòng 2-6: Dữ liệu (Lưu ý: NV01 có 2 thiết bị, NV02 có 3 thiết bị)
    data = [
        ["NV01", "Nguyễn Văn A", "01/01/1990", "15000000", 1, "Laptop Dell", 1],
        ["NV01", "Nguyễn Văn A", "01/01/1990", "15000000", 2, "Chuột không dây", 1],
        ["NV02", "Trần Thị B", "15/05/1995", "20000000", 1, "Macbook Pro", 1],
        ["NV02", "Trần Thị B", "15/05/1995", "20000000", 2, "Bàn phím cơ", 1],
        ["NV02", "Trần Thị B", "15/05/1995", "20000000", 3, "Màn hình rời", 2],
    ]

    for row in data:
        ws.append(row)

    try:
        wb.save("samples/Du_Lieu_Nhan_Su.xlsx")
        print("✅ Đã tạo file: samples/Du_Lieu_Nhan_Su.xlsx")
    except PermissionError:
        print(
            "❌ LỖI: Không thể lưu. Vui lòng đóng file 'samples/Du_Lieu_Nhan_Su.xlsx' nếu đang mở trong Excel và chạy lại."
        )


if __name__ == "__main__":
    generate_word()
    generate_excel()
    print("🎉 Hoàn tất! Mời bạn vào thư mục 'samples' để lấy file.")
