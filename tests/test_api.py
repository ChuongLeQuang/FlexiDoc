"""
EN: Unit tests for FastAPI endpoints using TestClient.
VI: Kịch bản kiểm thử cho các endpoint API sử dụng TestClient.
"""

import os
import pytest
import io
import time
import json
import zipfile
import docx
import openpyxl
from fastapi.testclient import TestClient
from src.main import app
from src.database.db import init_db, SQLALCHEMY_DATABASE_URL


@pytest.fixture(scope="module", autouse=True)
def setup_teardown_db():
    # Xóa file DB cũ trước khi test để đảm bảo môi trường sạch
    db_file = SQLALCHEMY_DATABASE_URL.split("///")[-1]
    if os.path.exists(db_file):
        os.remove(db_file)
    init_db()
    yield


client = TestClient(app)


def create_dummy_docx() -> bytes:
    doc = docx.Document()
    doc.add_paragraph("Họ tên: {{ ho_ten }}")
    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


def create_dummy_excel() -> bytes:
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Sheet1"
    ws.append(["Họ và tên", "Tuổi"])
    ws.append(["Nguyễn Văn A", 30])
    out = io.BytesIO()
    wb.save(out)
    return out.getvalue()


def test_validate_template_api():
    file_bytes = create_dummy_docx()
    response = client.post(
        "/api/v1/templates/validate",
        files={
            "file": (
                "test.docx",
                file_bytes,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    assert response.status_code == 200
    assert "ho_ten" in response.json()["single_vars"]


def test_validate_excel_api():
    file_bytes = create_dummy_excel()
    response = client.post(
        "/api/v1/excel/validate",
        data={"sheet_name": "Sheet1", "header_row": 1},
        files={
            "file": (
                "test.xlsx",
                file_bytes,
                "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            )
        },
    )
    assert response.status_code == 200
    assert "Họ và tên" in response.json()["columns"]


def test_profiles_crud_api():
    # Dùng timestamp để tạo tên unique tránh lỗi trùng lặp khi chạy test nhiều lần
    unique_name = f"Test Profile {time.time()}"
    payload = {
        "name": unique_name,
        "template_filename": "mau_hop_dong.docx",
        "sheet_name": "Data",
        "header_row": 1,
        "rules": [],
    }

    res_post = client.post("/api/v1/profiles/", json=payload)
    assert res_post.status_code == 200
    assert res_post.json()["name"] == unique_name

    res_get = client.get("/api/v1/profiles/")
    assert res_get.status_code == 200
    assert isinstance(res_get.json(), list)


def test_generate_documents_api():
    """Kiểm thử API trộn và sinh file ZIP (Tạo hợp đồng/Phụ lục)."""
    word_bytes = create_dummy_docx()
    excel_bytes = create_dummy_excel()

    rules = [
        {
            "variable_type": "single",
            "word_var": "ho_ten",
            "excel_col": "Họ và tên",
            "format_type": "text",
        }
    ]

    data = {
        "sheet_name": "Sheet1",
        "header_row": 1,
        "grouping_key": "",
        "filename_pattern": "HopDong_{{ ho_ten }}",
        "rules": json.dumps(rules),
    }

    files = {
        "word_file": (
            "template.docx",
            word_bytes,
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ),
        "excel_file": (
            "data.xlsx",
            excel_bytes,
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        ),
    }

    response = client.post("/api/v1/documents/generate", data=data, files=files)

    assert response.status_code == 200
    assert response.headers["content-type"] == "application/zip"

    # Đọc thử file ZIP trả về trên RAM để xác nhận có chứa file Word đã render
    with zipfile.ZipFile(io.BytesIO(response.content), "r") as zip_ref:
        file_names = zip_ref.namelist()
        assert len(file_names) > 0
        # Do dòng dữ liệu Excel là "Nguyễn Văn A", tên file ZIP sinh ra phải chứa "HopDong"
        assert any("HopDong" in name for name in file_names)
        assert any(name.endswith(".docx") for name in file_names)


def test_validate_template_zero_variables_api():
    """
    EN: Test API response when Word file has no variables (ZeroVariableError).
    VI: Test API bẫy lỗi ZeroVariableError khi file Word không có biến.
    """
    doc = docx.Document()
    doc.add_paragraph(
        "Đây là một file Word bình thường, không chứa thẻ ngoặc nhọn nào."
    )
    out = io.BytesIO()
    doc.save(out)

    response = client.post(
        "/api/v1/templates/validate",
        files={
            "file": (
                "test_no_var.docx",
                out.getvalue(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    # Tuỳ thuộc vào cấu hình Exception Handler ở main.py, thường lỗi nghiệp vụ sẽ là 400
    assert response.status_code >= 400
    assert response.json().get("error_code") == "ZERO_VARIABLES"


def test_validate_template_invalid_format_api():
    """Test API phản hồi lỗi khi file tải lên bị hỏng hoặc không đúng định dạng."""
    response = client.post(
        "/api/v1/templates/validate",
        files={"file": ("bad_file.docx", b"Garbage data", "application/msword")},
    )
    assert response.status_code >= 400


def test_validate_excel_sheet_not_found_api():
    """Test API bẫy lỗi SheetNotFoundError khi chọn sai tên Sheet."""
    file_bytes = create_dummy_excel()  # Hàm này tạo file có sheet tên là "Sheet1"

    response = client.post(
        "/api/v1/excel/validate",
        data={"sheet_name": "SheetKhongTonTai", "header_row": 1},
        files={
            "file": (
                "test.xlsx",
                file_bytes,
                "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            )
        },
    )

    assert response.status_code >= 400
    assert "SheetKhongTonTai" in response.json().get("message", "")
