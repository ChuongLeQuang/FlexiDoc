"""
EN: Unit tests for Word Service.
VI: Kịch bản kiểm thử cho Dịch vụ Word.
"""

import pytest
import io
import docx
from src.services.word_service import (
    extract_variables,
    preprocess_template,
    validate_syntax,
)
from src.exceptions.custom_errors import (
    InvalidFileFormatError,
    ZeroVariableError,
    TemplateSyntaxError,
)


def create_dummy_docx(text: str) -> bytes:
    """Tạo một file Word ảo trên RAM để test."""
    doc = docx.Document()
    doc.add_paragraph(text)
    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


def test_extract_variables_zero():
    """Test lỗi khi không có biến nào."""
    template = create_dummy_docx("Đây là file không có biến.")
    with pytest.raises(ZeroVariableError):
        extract_variables(template)


def test_extract_variables_success():
    """Test trích xuất biến đơn và biến bảng thành công."""
    template = create_dummy_docx(
        "Họ tên: {{ ho_ten }}\n"
        "{{ table_start:danh_sach }}\n"
        "STT: {{ stt }} - Lương: {{ item.luong }}\n"
        "{{ table_end }}"
        "\n{{ list_start:danh_sach_lop }}\n"
        "Lớp: {{ item.ten_lop }}\n"
        "{{ list_end }}"
    )
    result = extract_variables(template)
    assert "ho_ten" in result["single_vars"]
    assert "danh_sach" in result["table_vars"]
    assert "stt" in result["table_vars"]["danh_sach"]
    assert "luong" in result["table_vars"]["danh_sach"]
    assert "danh_sach_lop" in result["table_vars"]
    assert "ten_lop" in result["table_vars"]["danh_sach_lop"]


def test_validate_syntax_error():
    """Test bắt lỗi sai cú pháp thẻ Jinja2 (thiếu dấu ngoặc)."""
    # Cố tình viết sai cú pháp (thiếu % ở cuối)
    # Tuy nhiên, docxtpl sử dụng {{ }} nên ta mô phỏng một thẻ if bị sai
    template = create_dummy_docx("Lương là: {% if luong > 1000 %} Cao {% endif")

    with pytest.raises(TemplateSyntaxError):
        validate_syntax(template)


def test_extract_variables_order_preservation():
    """
    EN: Test if extract_variables preserves the exact order of appearance.
    VI: Kiểm tra xem hàm extract_variables có bảo toàn đúng thứ tự xuất hiện không.
    """
    text = (
        "Xin chào {{ ten_khach_hang }}, "
        "Hợp đồng số {{ so_hop_dong }} được lập ngày {{ ngay_lap }}. "
        "{{ table_start:ds_san_pham }} "
        "{{ item.stt }} | {{ item.ma_sp }} | {{ item.ten_sp }} "
        "{{ table_end }}"
    )
    template = create_dummy_docx(text)
    result = extract_variables(template)

    assert result["single_vars"] == ["ten_khach_hang", "so_hop_dong", "ngay_lap"]
    assert "ds_san_pham" in result["table_vars"]
    assert result["table_vars"]["ds_san_pham"] == ["stt", "ma_sp", "ten_sp"]


def test_extract_variables_duplicates_order():
    """
    EN: Test if duplicates are ignored while keeping the first appearance order.
    VI: Kiểm tra xử lý biến trùng lặp và bảo toàn thứ tự xuất hiện lần đầu.
    """
    text = (
        "{{ b_var }} {{ a_var }} {{ b_var }} {{ c_var }} {{ a_var }} "
        "{{ table_start:tb }} {{ item.y }} {{ item.x }} {{ item.y }} {{ table_end }}"
    )
    template = create_dummy_docx(text)
    result = extract_variables(template)

    assert result["single_vars"] == ["b_var", "a_var", "c_var"]
    assert result["table_vars"]["tb"] == ["y", "x"]


def test_invalid_file_format_error():
    """
    EN: Test error when file format is invalid/corrupted.
    VI: Test lỗi khi truyền vào file Word hỏng hoặc không đúng định dạng docx.
    """
    # Truyền chuỗi byte rác không phải là cấu trúc zip/docx hợp lệ
    bad_bytes = b"This is just a random text string, not a real word document."
    with pytest.raises(InvalidFileFormatError):
        extract_variables(bad_bytes)


def test_preprocess_template_list_success():
    """
    EN: Test successful conversion of list tags to Jinja2 syntax in paragraphs.
    VI: Test tự động chuyển đổi thẻ danh sách đơn giản sang cú pháp Jinja2 trong đoạn văn.
    """
    template = create_dummy_docx("{{ list_start:ds_nhan_vien }} {{ list_end }}")
    result_bytes = preprocess_template(template)

    processed_doc = docx.Document(io.BytesIO(result_bytes))
    combined_text = "\n".join([p.text for p in processed_doc.paragraphs])

    assert "{%p for item in ds_nhan_vien %}" in combined_text
    assert "{%p endfor %}" in combined_text


def test_preprocess_template_table_success():
    """
    EN: Test successful conversion of table tags to Jinja2 syntax in table rows.
    VI: Test tự động chuyển đổi thẻ bảng biểu đơn giản sang cú pháp Jinja2 trong bảng.
    """
    doc = docx.Document()
    table = doc.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "{{ table_start:ds_thiet_bi }} {{ table_end }}"
    out = io.BytesIO()
    doc.save(out)

    result_bytes = preprocess_template(out.getvalue())

    processed_doc = docx.Document(io.BytesIO(result_bytes))
    cell_text = processed_doc.tables[0].cell(0, 0).text

    assert "{%tr for item in ds_thiet_bi %}" in cell_text
    assert "{%tr endfor %}" in cell_text
